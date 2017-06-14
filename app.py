import argparse, requests, json
from docx import Document
from docx.shared import RGBColor
import datetime

THRESH_VAL = 0.75
RED = RGBColor(252, 42, 53)
LAVENDER = RGBColor(100, 98, 150)

def request(apikey, trans_id):
    r = requests.get(
        'https://api.capio.ai/v1/speech/transcript/%s' % trans_id,
        headers={
            'apiKey': apikey,
        }
    )
    if r.status_code != 200:
        raise Exception('Request returned a %d status code' % r.status_code)
    return r.content
    
def format_time(s):
    temp = round(s, 2)
    #print int(round(s,2)*100%100*1000000)
    t = datetime.datetime(2009, 2, 2, int(s)/3600, int(s)/60, int(s)%60, int(round(s,2)*100%100*10000))
    return t.strftime('%H:%M:%S.%f')[:-4]
    
class MS_Doc(object):
    def __init__(self, debug=False):
        parser = self._parser()
        self.args = parser.parse_args()
        
    def _parser(self):
        parser = argparse.ArgumentParser()
        return parser
    
    def format_line(self, line):
        return line
    
    def write_doc(self, apikey, trans_id, file_name='output.docx'):
        response = request(
                    apikey=apikey,
                    trans_id=trans_id
                )
        doc = Document()
        parsed_data = json.loads(response)
        for line in parsed_data:
            words = line['result'][0]['alternative'][0]['words']
            p = doc.add_paragraph(format_time(words[0]['from']))
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = LAVENDER
            p.add_run('\t')
            for i, word in enumerate(words):
                if i > 0:
                    p.add_run(' ')
                run = p.add_run(word['word'])
                if word['confidence'] <= THRESH_VAL:
                    font = run.font
                    font.color.rgb = RED

        doc.save(file_name)

MS_Doc().write_doc('262ac9a0c9ba4d179aad4c0b9b02120a', '593f237fbcae700012ba8fcd')
    